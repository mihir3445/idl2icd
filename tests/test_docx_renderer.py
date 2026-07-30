
import pytest

try:
    from docx import Document
except ImportError:
    pytest.skip("python-docx is not installed", allow_module_level=True)

from idl2icd.model.ir import ProjectMeta
from idl2icd.model.merge import build_ir
from idl2icd.render.docx.renderer import render_docx

IDL_TEXT = """
@topic
struct Foo {
  @key unsigned long id;
  float value;
};
"""
META_TEXT = """
topics:
  Foo:
    description: "A test topic."
    criticality: high
    qos:
      overrides: { reliability: RELIABLE, durability: TRANSIENT_LOCAL }
    publishers:
      - participant: A
    subscribers:
      - participant: B
    fields:
      value:
        unit: percent
"""


def test_render_docx_produces_valid_document_with_expected_content(tmp_path):
    idl = tmp_path / "x.idl"
    idl.write_text(IDL_TEXT)
    meta = tmp_path / "m.yaml"
    meta.write_text(META_TEXT)

    project = ProjectMeta(name="Test ICD", version="1.0.0", organization="M")
    ir = build_ir([idl], [meta], project)

    out_path = tmp_path / "out.docx"
    render_docx(ir, diagnostics=[], out_path=out_path)

    assert out_path.exists()

    doc = Document(str(out_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    all_text += "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )

    assert "Test ICD" in all_text
    assert "Foo" in all_text
    assert "A test topic." in all_text
    assert "value" in all_text
    assert "percent" in all_text
    assert any(t.rows for t in doc.tables)  # at least one populated table
